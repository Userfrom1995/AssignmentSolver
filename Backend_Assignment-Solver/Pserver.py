from flask import Flask, send_from_directory, jsonify, request
from flask_cors import CORS
import os
from PIL import Image
import io
import PyPDF2
from paddleocr import PaddleOCR  # Use PaddleOCR for OCR
from docx import Document
import fitz
import numpy as np
import google.generativeai as genai
from datetime import datetime

app = Flask(__name__)
CORS(app, resources={r"/*": {"origins": "http://localhost:4200"}})
genai.configure(api_key=os.getenv("api_key"))
UPLOAD_FOLDER = "Questions"  # Specify the path to save uploaded files
SOLUTION_FOLDER = "Solutions"

if not os.path.exists(UPLOAD_FOLDER):
    os.makedirs(UPLOAD_FOLDER)

if not os.path.exists(SOLUTION_FOLDER):
    os.makedirs(SOLUTION_FOLDER)

# Initialize PaddleOCR
ocr = PaddleOCR(use_angle_cls=True, lang='en', gpu=True)

@app.route('/solve-assignment', methods=['POST'])
def solve_assignment():
    file = request.files['file']
    data = request.form.to_dict()
    academic_level = data.get('academicLevel', 'college-level')
    subject = data.get('subject', 'General')
    additional_info = data.get('additionalInfo', '')

    # Save the file
    file_path = os.path.join(UPLOAD_FOLDER, file.filename)
    file.save(file_path)

    # Process the file (OCR)
    if file.filename.lower().endswith(('.jpg', '.jpeg', '.png')):
        extracted_text = ocr_image(file_path)
    elif file.filename.lower().endswith('.pdf'):
        extracted_text = ocr_pdf(file_path)
    else:
        return jsonify({'error': 'Unsupported file format. Please provide an image (JPEG/PNG) or PDF file.'})

    # Generate a response (you can add your logic here)
    prompt = f"""You are a helpful and intelligent AI assistant designed to solve academic assignments for {academic_level} students. You will receive text extracted from images or PDF documents, which will contain the assignment questions.

    Your task is to:

    1. Carefully analyze the provided text.
    2. Identify and extract all assignment questions. Questions can be phrased in various ways (e.g., direct questions, problem statements, instructions to solve).
    3. Solve each question accurately and comprehensively, keeping in mind that the subject is {subject}. Show your work, providing step-by-step solutions, explanations, code (if applicable), and any relevant diagrams or figures.
    4. Format your responses in a clear, organized, and human-readable manner. Use proper headings, numbering, and formatting conventions as appropriate for the assignment type.

    Additional Instructions:

    - If the assignment requires specific software or tools, you can mention them and provide general instructions but cannot directly execute code or commands.
    - If you encounter any unclear or ambiguous questions, state your assumptions clearly and provide the best possible solution based on your understanding.
    - Maintain a professional and academic tone throughout your responses.
    - {additional_info}

    Input Format:
          A single string containing the extracted text from the image or PDF."""
          
    # Placeholder
    model = genai.GenerativeModel('gemini-1.5-flash')
    chat = model.start_chat(history=[])
    response = chat.send_message(prompt)
    response = chat.send_message(extracted_text)

    # Save the response as a Word document
    solution_file_name = file.filename.replace(".pdf", ".docx").replace(".png", ".docx").replace(".jpg", ".docx").replace(".jpeg", ".docx")
    solution_file_path = os.path.join(SOLUTION_FOLDER, solution_file_name)
    save_to_word(response, solution_file_path)

    return jsonify({'message': 'Assignment solved successfully', 'output_path': solution_file_path})

def ocr_image(file_path):
    # Use PaddleOCR to extract text from image
    result = ocr.ocr(file_path, cls=True)
    extracted_text = ''
    for line in result:
        for word_info in line:
            extracted_text += word_info[1][0] + ' '
    return extracted_text.strip()

def ocr_pdf(file_path):
    # Open the PDF
    pdf_document = fitz.open(file_path)
    extracted_text = ""

    # Loop through each page
    for page_num in range(len(pdf_document)):
        page = pdf_document[page_num]
        
        # Extract text directly from the page
        text = page.get_text()
        if text:
            extracted_text += text
        else:
            # If no text is found, try to extract images
            image_list = page.get_images(full=True)
            for img_index, img in enumerate(image_list):
                xref = img[0]
                base_image = pdf_document.extract_image(xref)
                image_bytes = base_image["image"]

                # Convert bytes to an image and run OCR
                img = Image.open(io.BytesIO(image_bytes))
                result = ocr.ocr(np.array(img))
                for line in result:
                    extracted_text += " ".join([word_info[1][0] for word_info in line])

    pdf_document.close()
    return extracted_text

def save_to_word(text, output_file_path):
    doc = Document()
    doc.add_paragraph(text.text)
    for para in doc.paragraphs:
        if '**' in para.text:
            para.text = para.text.replace('**', '')
    doc.save(output_file_path)

@app.route('/api/files', methods=['GET'])
def list_files():
    files = os.listdir(SOLUTION_FOLDER)
    now = datetime.now()
    formatted_date_time = now.strftime("%Y-%m-%d %H:%M:%S")
    file_list = [
        {"name": file, "size": os.path.getsize(os.path.join(SOLUTION_FOLDER, file)), "date": formatted_date_time
         } for file in files]
    return jsonify(file_list)

@app.route('/api/files/<filename>', methods=['GET'])
def download_file(filename):
    return send_from_directory(SOLUTION_FOLDER, filename)


@app.route('/api/files/<filename>', methods=['DELETE'])
def delete_file(filename):
    os.remove(os.path.join(SOLUTION_FOLDER, filename))
    return '', 204

if __name__ == "__main__":
    app.run(debug=True)
