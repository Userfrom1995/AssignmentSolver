# Import the required libraries
from openai import OpenAI
import pytesseract
from PIL import Image
import io
import PyPDF2
import json
from bs4 import BeautifulSoup
from reportlab.lib.pagesizes import letter
from reportlab.pdfgen import canvas
from docx import Document
import warnings
from collections.abc import Sequence


warnings.simplefilter("ignore", PendingDeprecationWarning)


# Path to the Tesseract executable (change this based on your installation)
pytesseract.pytesseract.tesseract_cmd = r"C:\Program Files\Tesseract-OCR\tesseract.exe"

# Initialize the OpenAI client
client = OpenAI(base_url="http://localhost:1234/v1", api_key="lm-studio")

def ocr_image(image):
    # Open the image file using PIL (Python Imaging Library)
    img = Image.open(io.BytesIO(image))

    # Perform OCR on the image using pytesseract
    extracted_text = pytesseract.image_to_string(img)

    return extracted_text

class MyClass:
    def __init__(self, value):
        self.value = value
    
    
    # def __str__(self):
    #     return f"{self.value})"



def save_to_pdf(text, output_file_path):
    # Create a canvas
    c = canvas.Canvas(output_file_path, pagesize=letter)

    # Set font and size
    c.setFont("Helvetica", 12)

    # Write text to PDF
    c.drawString(100, 700, text)

    # Save the PDF file
    c.save()

# def save_to_word(text, output_file_path):
#     # Create a new Word document
#     doc = Document()
#
#     # Add a paragraph with the text to the document
#     doc.add_paragraph(text)
#
#     # Save the Word document
#     doc.save(output_file_path)

def save_to_word_from_chat_completion(chat_completion, output_file_path):
    # Create a new Word document
    doc = Document()

    # Extract content from the ChatCompletionMessage
    content = chat_completion.choices[0].message.content

    # Add a paragraph with the extracted content to the document
    doc.add_paragraph(content)

    # Save the Word document
    doc.save(output_file_path)

def ocr_pdf(pdf):
    # Open the PDF file using PyPDF2
    pdf_reader = PyPDF2.PdfReader(io.BytesIO(pdf))
    num_pages = len(pdf_reader.pages)

    extracted_text = ""

    # Extract text from each page of the PDF
    for page_num in range(num_pages):
        page = pdf_reader.pages[page_num]
        extracted_text += page.extract_text()

    return extracted_text

def classify_and_correct_text(text):
    # Pass the text to the local AI model (Lama) for classification and correction
    completion = client.chat.completions.create(
        model="lmstudio-community/Meta-Llama-3-8B-Instruct-GGUF",
        messages=[
            {
                "role": "system",
                   "content": "solve the question in well structured way"   },
            {
                "role": "user",
                "content": [
                    {"type": "text", "text": text},
                ],
            }
        ],
        max_tokens=1000,
        stream=False,
        timeout=600
    )

    # Extract and return the corrected text from the response
    # for chunk in completion:
    #     if chunk.choices is not None and chunk.choices[0].delta.content:
    #         return print(chunk.choices[0].delta.content, end="", flush=True)
    #     else:
    #        return print("Received chunk:", chunk)
    # for chunk in completion:
    #     if chunk[0]:  # Check if response is complete
    #         # Extract the content from the last message in the list
    #         messages = chunk[1]
    #         last_message = messages
    #         content = last_message
    #         print(content)
    #     else:
    #         print("Received incomplete response. Waiting for more chunks...")

    # corrected_text = completion['choices'][0]['message']['content'][0]['text']
    # return corrected_text
    # Create an instance of the class
    # obj = MyClass(completion)
    # string = str(obj)
    

    # Print the response
    return completion
     # for chunk in completion:
     #         if chunk.choices[0].delta.content:
     #           print(chunk.choices[0].delta.content, end="", flush=True)
     #           new_message["content"] += chunk.choices[0].delta.content)
     #
    


def replace_slash_n(text):
    # Replace all occurrences of '\n' with actual newline characters '\n'
    replaced_text = text.replace('\\n', '\n')
    return replaced_text

# Example usage
if __name__ == "__main__":
    # Example input file path
    input_file_path = input("Enter the path to the input file (image or PDF): ")

    # Read input file
    with open(input_file_path, 'rb') as file:
        file_content = file.read()

    # Perform OCR based on file type
    if input_file_path.lower().endswith(('.jpg', '.jpeg', '.png')):
        extracted_text = ocr_image(file_content)
    elif input_file_path.lower().endswith('.pdf'):
        extracted_text = ocr_pdf(file_content)
    else:
        print("Unsupported file format. Please provide an image (JPEG/PNG) or PDF file.")

    # Print the extracted text
    print("Extracted Text:")
    print(extracted_text)

    # Classify and correct the text
    corrected_text1 = (
    classify_and_correct_text(extracted_text))
    print(corrected_text1)

    # save_to_pdf(corrected_text1, "Answers/output.pdf")
    save_to_word_from_chat_completion(corrected_text1, "Answers/output.docx")

    # corrected_text1 = corrected_text.replace('\\n', '\n')
    # Print the corrected text
    # print("\nCorrected Text:")
    # print(corrected_text)
    # replaced_text = replace_slash_n(corrected_text)
    # print("Replaced Text:")
    # print(corrected_text)



    # Assuming response is the variable containing the response object
    # corrected_text1 = corrected_text1.json()
    # print(corrected_text1)# Convert response to JSON format

    # Assuming response_json contains the JSON-formatted response
    # soup = BeautifulSoup(corrected_text1[0], 'html.parser')  # Parse the HTML content using Beautiful Soup
    #
    # questions = []  # List to store extracted questions
    #
    # # Extract questions from the parsed HTML content
    # for question in soup.find_all('b', string=lambda t: t and t.startswith('**Question')):
    #     question_text = question.next_sibling.strip()  # Get the text of the question
    #     questions.append(question_text)
    #     print(question)

